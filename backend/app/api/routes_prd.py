import json
import re
from pathlib import Path
from fastapi import APIRouter, Depends
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from app.agents.metric_analyst_agent import analyze_metrics
from app.agents.prd_writer_agent import generate_prd
from app.agents.reviewer_agent import review_prd
from app.db.database import get_db
from app.db.models import PrdDocument
from app.core.config import get_settings
from app.schemas.prd import PrdExportRequest, PrdGenerateRequest, PrdUpdateRequest

router = APIRouter(prefix="/api/prd", tags=["prd"])


@router.post("/generate")
async def generate(payload: PrdGenerateRequest, db: Session = Depends(get_db)):
    prd = await generate_prd(db, payload.opportunity_id, payload.project_id, analyze_metrics(db, payload.project_id, payload.conversation_id), payload.conversation_id)
    return serialize(prd)


@router.get("/{prd_id}")
def get_prd(prd_id: int, db: Session = Depends(get_db)):
    return serialize(db.get(PrdDocument, prd_id))


@router.get("")
def list_prds(conversation_id: str | None = None, db: Session = Depends(get_db)):
    q = db.query(PrdDocument)
    if conversation_id:
        q = q.filter(PrdDocument.conversation_id == conversation_id)
    return [serialize(p) for p in q.order_by(PrdDocument.id.desc()).all()]


@router.post("/{prd_id}/review")
async def review(prd_id: int, db: Session = Depends(get_db)):
    return await review_prd(db, prd_id)


@router.post("/{prd_id}/update")
def update_prd(prd_id: int, payload: PrdUpdateRequest, db: Session = Depends(get_db)):
    prd = db.get(PrdDocument, prd_id)
    prd.prd_markdown = payload.prd_markdown
    db.commit()
    db.refresh(prd)
    return serialize(prd)


@router.post("/export-docx")
def export_docx(payload: PrdExportRequest):
    from docx import Document

    settings = get_settings()
    safe_title = re.sub(r"[^A-Za-z0-9_\-\u4e00-\u9fff]+", "_", payload.title).strip("_") or "prd"
    path = settings.export_dir / f"{safe_title}.docx"
    doc = Document()
    for line in payload.prd_markdown.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line).strip(), style="List Number")
        elif line.strip():
            doc.add_paragraph(line.strip())
        else:
            doc.add_paragraph("")
    doc.save(path)
    return FileResponse(path, filename=path.name, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


def serialize(p: PrdDocument):
    return {"id": p.id, "project_id": p.project_id, "conversation_id": p.conversation_id, "opportunity_id": p.opportunity_id, "title": p.title, "version": p.version, "status": p.status, "prd_markdown": p.prd_markdown, "created_at": p.created_at.isoformat(), "updated_at": p.updated_at.isoformat()}
